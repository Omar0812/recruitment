"""简历内容提取：按文件格式分治路由。

返回统一结构：
  images 类型: {"type": "images", "content": [base64_str, ...]}
  text 类型:   {"type": "text", "content": "合并文本"}
  错误:        {"error": "...", "error_type": "extract_failed"}
"""
from __future__ import annotations

import base64
import logging
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_content(file_path: str) -> dict:
    """按文件后缀分治路由，返回统一结构。"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _extract_pdf(file_path)
        elif suffix == ".docx":
            return {"type": "text", "content": _extract_docx(file_path)}
        elif suffix in (".png", ".jpg", ".jpeg"):
            return {"type": "images", "content": [_extract_image(file_path)]}
        else:
            return {"error": f"不支持的文件类型：{suffix}", "error_type": "extract_failed"}
    except Exception as e:
        return {"error": str(e), "error_type": "extract_failed"}


def _extract_pdf(file_path: str) -> dict:
    """LiteParse 双路径提取 PDF：文本优先，截图 fallback。

    返回完整 dict（含 type 字段），由 extract_content 直接透传。
    """
    # 1. 尝试文本提取
    text = _liteparse_text(file_path)
    if text and len("".join(text.split())) >= 100:
        logger.info("PDF 走文本路径: %s", file_path)
        return {"type": "text", "content": text}

    # 2. 文本不足（设计师简历/扫描件），fallback 到截图
    logger.info("PDF 文本不足，fallback 到截图路径: %s", file_path)
    images = _liteparse_screenshots(file_path)
    if images:
        return {"type": "images", "content": images}

    # 3. 都失败
    raise RuntimeError("LiteParse 无法提取 PDF 内容（文本提取和截图均失败）")


def _liteparse_text(file_path: str) -> str | None:
    """LiteParse CLI 提取 PDF 文本（不启用 OCR）。

    成功返回文本内容，失败返回 None。
    """
    try:
        result = subprocess.run(
            ["lit", "parse", file_path, "--format", "text", "--no-ocr", "-q"],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        return None
    except Exception:
        logger.debug("LiteParse text 提取异常: %s", file_path, exc_info=True)
        return None


def _liteparse_screenshots(file_path: str) -> list[str] | None:
    """LiteParse CLI 生成 PDF 页面截图（前 10 页，150 DPI）。

    成功返回 base64 编码的 PNG 列表，失败返回 None。
    截图文件名格式为 page_1.png, page_2.png, ...（按文件名排序保证页序）。
    """
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [
                    "lit", "screenshot", file_path,
                    "--dpi", "150",
                    "--target-pages", "1-10",
                    "-o", tmpdir,
                    "-q",
                ],
                capture_output=True,
                timeout=60,
            )
            if result.returncode != 0:
                return None

            # 按文件名排序读取所有 .png 文件
            png_files = sorted(
                p for p in Path(tmpdir).iterdir() if p.suffix.lower() == ".png"
            )
            if not png_files:
                return None

            return [base64.b64encode(p.read_bytes()).decode() for p in png_files]
    except Exception:
        logger.debug("LiteParse screenshot 异常: %s", file_path, exc_info=True)
        return None


def _extract_docx(file_path: str) -> str:
    """python-docx 提取段落文本 + 表格内容，返回合并文本。"""
    from docx import Document

    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_image(file_path: str) -> str:
    """读取原始图片文件，返回 base64 编码。"""
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode()

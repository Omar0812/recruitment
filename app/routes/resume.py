import uuid
import shutil
from pathlib import Path
from html import escape
from fastapi import APIRouter, UploadFile, File, HTTPException
from app.extractor import extract_text
from app import ai_client

router = APIRouter(prefix="/api/resume", tags=["resume"])

RESUME_DIR = Path("data/resumes")


@router.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    raw_name = (file.filename or "").strip()
    safe_name = Path(raw_name.replace("\\", "/")).name
    suffix = Path(safe_name).suffix.lower()
    if suffix not in (".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"):
        raise HTTPException(status_code=400, detail="不支持的文件格式")
    if not safe_name:
        raise HTTPException(status_code=400, detail="文件名无效")

    file_id = str(uuid.uuid4())
    dest_dir = RESUME_DIR / file_id
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / safe_name

    with open(dest_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    text = extract_text(str(dest_path))

    ai_configured = ai_client._is_configured()
    parsed = {}
    warning = None

    if ai_configured:
        try:
            parsed = ai_client.parse_resume(text)
        except Exception as e:
            warning = f"AI解析失败: {e}"
    else:
        warning = "AI未配置，请在 config.json 中填入 API Key 后重启应用"

    # Build static preview URL
    preview_url = f"/resumes/{file_id}/{safe_name}"
    preview_type = "file"
    preview_html = None

    if suffix == ".pdf":
        preview_type = "pdf"
    elif suffix in (".jpg", ".jpeg", ".png"):
        preview_type = "image"
    elif suffix == ".docx":
        preview_type = "word"
        try:
            from docx import Document
            doc = Document(str(dest_path))
            html_parts = []
            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt:
                    html_parts.append(f"<p>{escape(txt)}</p>")
            for table in doc.tables:
                html_parts.append("<table style='border-collapse:collapse;width:100%'>")
                for row in table.rows:
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        html_parts.append(
                            "<td style='border:1px solid #ddd;padding:6px;vertical-align:top'>"
                            f"{escape(cell.text.strip())}</td>"
                        )
                    html_parts.append("</tr>")
                html_parts.append("</table>")
            preview_html = "".join(html_parts) or "<p>Word 内容为空</p>"
        except Exception:
            preview_html = f"<pre>{escape(text or '无法解析 Word 内容')}</pre>"
    elif suffix == ".doc":
        preview_type = "word"
        preview_html = f"<pre>{escape(text or '无法解析 Word 内容')}</pre>"

    return {
        "file_id": file_id,
        "resume_path": str(dest_path),
        "file_suffix": suffix,
        "preview_url": preview_url,
        "preview_type": preview_type,
        "preview_html": preview_html,
        "extracted_text": text,
        "parsed": parsed,
        "warning": warning,
    }

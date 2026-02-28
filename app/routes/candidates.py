from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.orm import Session
from sqlalchemy import or_, func
from pydantic import BaseModel
from typing import Optional, List, Literal
from datetime import datetime
import os
from html import escape

from app.database import get_db
from app.models import Candidate, HistoryEntry, Supplier, CandidateJobLink, ActivityRecord
from app.schemas import CandidateOut
from app.services import candidates as candidates_svc

router = APIRouter(prefix="/api/candidates", tags=["candidates"])


def _normalize_tags(tags: Optional[List[str]]) -> List[str]:
    if not tags:
        return []
    seen = set()
    result = []
    for item in tags:
        t = str(item or "").strip()
        if not t:
            continue
        key = t.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(t)
    return result


def _parse_tag_query(tags: Optional[str]) -> List[str]:
    if not tags:
        return []
    parts = [x.strip() for x in tags.replace("，", ",").split(",")]
    return _normalize_tags(parts)


class CandidateCreate(BaseModel):
    name: Optional[str] = None
    name_en: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    age: Optional[int] = None
    education: Optional[str] = None
    school: Optional[str] = None
    city: Optional[str] = None
    last_company: Optional[str] = None
    last_title: Optional[str] = None
    years_exp: Optional[float] = None
    skill_tags: Optional[List[str]] = []
    source: Optional[str] = None
    referred_by: Optional[str] = None
    supplier_id: Optional[int] = None
    notes: Optional[str] = None
    resume_path: Optional[str] = None
    education_list: Optional[List[dict]] = []
    work_experience: Optional[List[dict]] = []
    project_experience: Optional[List[dict]] = []


class CandidateUpdate(BaseModel):
    name: Optional[str] = None
    name_en: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    age: Optional[int] = None
    education: Optional[str] = None
    school: Optional[str] = None
    city: Optional[str] = None
    last_company: Optional[str] = None
    last_title: Optional[str] = None
    years_exp: Optional[float] = None
    skill_tags: Optional[List[str]] = None
    source: Optional[str] = None
    referred_by: Optional[str] = None
    supplier_id: Optional[int] = None
    notes: Optional[str] = None
    followup_status: Optional[str] = None
    education_list: Optional[List[dict]] = None
    work_experience: Optional[List[dict]] = None
    project_experience: Optional[List[dict]] = None
    starred: Optional[bool] = None


def _candidate_out(c: Candidate) -> dict:
    return CandidateOut.model_validate(c).model_dump()


def _sync_legacy_fields(candidate: Candidate):
    """从新数组字段同步旧字段（向后兼容）"""
    edu_list = candidate.education_list or []
    if edu_list:
        candidate.education = edu_list[0].get("degree") or candidate.education
        candidate.school = edu_list[0].get("school") or candidate.school
    work_list = candidate.work_experience or []
    if work_list:
        candidate.last_company = work_list[0].get("company") or candidate.last_company
        candidate.last_title = work_list[0].get("title") or candidate.last_title


def _normalize_text(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _has_value(value) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, list):
        return len(value) > 0
    return True


def _last_application_context(candidate_id: int, db: Session) -> Optional[dict]:
    link = (
        db.query(CandidateJobLink)
        .filter(CandidateJobLink.candidate_id == candidate_id)
        .order_by(CandidateJobLink.updated_at.desc(), CandidateJobLink.id.desc())
        .first()
    )
    if not link:
        return None

    last_interview = (
        db.query(ActivityRecord)
        .filter(ActivityRecord.link_id == link.id, ActivityRecord.type == "interview")
        .order_by(ActivityRecord.id.desc())
        .first()
    )
    summary = None
    if last_interview:
        payload = last_interview.payload or {}
        comment = payload.get("comment") or last_interview.comment or ""
        conclusion = payload.get("conclusion") or last_interview.conclusion or ""
        summary = f"{conclusion}：{comment[:50]}" if comment else conclusion

    days_ago = (datetime.utcnow() - link.updated_at).days if link.updated_at else None
    return {
        "job_title": link.job.title if link.job else None,
        "final_stage": link.stage,
        "outcome": link.outcome,
        "rejection_reason": link.rejection_reason,
        "days_ago": days_ago,
        "last_interview_summary": summary,
    }


def _active_link_context(candidate_id: int, db: Session) -> Optional[dict]:
    link = (
        db.query(CandidateJobLink)
        .filter(
            CandidateJobLink.candidate_id == candidate_id,
            CandidateJobLink.outcome.is_(None),
        )
        .order_by(CandidateJobLink.updated_at.desc(), CandidateJobLink.id.desc())
        .first()
    )
    if not link:
        return None
    return {
        "id": link.id,
        "job_id": link.job_id,
        "job_title": link.job.title if link.job else None,
        "stage": link.stage,
    }


def _create_candidate_record(payload: dict, db: Session) -> Candidate:
    if not payload.get("name") and payload.get("name_en"):
        payload["name"] = payload["name_en"]
    if not payload.get("name"):
        raise HTTPException(status_code=400, detail="姓名不能为空")

    candidate = Candidate(**payload)
    if candidate.supplier_id:
        supplier = db.query(Supplier).filter(Supplier.id == candidate.supplier_id).first()
        if supplier:
            candidate.source = supplier.name
    _sync_legacy_fields(candidate)
    db.add(candidate)
    db.flush()
    db.add(HistoryEntry(candidate_id=candidate.id, event_type="created", detail="候选人档案创建"))
    return candidate


def _merge_skill_tags(existing_tags: Optional[List[str]], incoming_tags: Optional[List[str]]) -> List[str]:
    merged = []
    for tag in (existing_tags or []) + (incoming_tags or []):
        normalized = _normalize_text(tag)
        if normalized and normalized not in merged:
            merged.append(normalized)
    return merged


def _merge_intake_payload(target: Candidate, incoming_payload: dict, overwrite_resume: bool, db: Session) -> List[str]:
    merged_fields = []

    if not _has_value(target.name) and _has_value(incoming_payload.get("name")):
        target.name = incoming_payload["name"]
        merged_fields.append("姓名")

    fill_if_empty_fields = [
        ("name_en", "英文名"),
        ("phone", "手机"),
        ("email", "邮箱"),
        ("age", "年龄"),
        ("education", "学历"),
        ("school", "学校"),
        ("city", "城市"),
        ("last_company", "当前公司"),
        ("last_title", "当前职位"),
        ("years_exp", "工作年限"),
        ("source", "来源"),
        ("referred_by", "内推人"),
        ("supplier_id", "供应商"),
    ]
    for field, label in fill_if_empty_fields:
        incoming_value = incoming_payload.get(field)
        current_value = getattr(target, field)
        if _has_value(incoming_value) and not _has_value(current_value):
            setattr(target, field, incoming_value)
            merged_fields.append(label)

    merged_tags = _merge_skill_tags(target.skill_tags, incoming_payload.get("skill_tags"))
    if merged_tags != (target.skill_tags or []):
        target.skill_tags = merged_tags
        merged_fields.append("技能标签")

    array_fields = [
        ("education_list", "教育经历"),
        ("work_experience", "工作经历"),
        ("project_experience", "项目经历"),
    ]
    for field, label in array_fields:
        incoming_array = incoming_payload.get(field) or []
        current_array = getattr(target, field) or []
        if incoming_array and not current_array:
            setattr(target, field, incoming_array)
            merged_fields.append(label)

    incoming_notes = _normalize_text(incoming_payload.get("notes"))
    if incoming_notes:
        current_notes = _normalize_text(target.notes)
        if not current_notes:
            target.notes = incoming_notes
            merged_fields.append("备注")
        elif incoming_notes not in current_notes:
            target.notes = f"{target.notes}\n\n[导入补充]\n{incoming_notes}"
            merged_fields.append("备注补充")

    incoming_resume_path = _normalize_text(incoming_payload.get("resume_path"))
    if incoming_resume_path and (overwrite_resume or not _normalize_text(target.resume_path)):
        if target.resume_path != incoming_resume_path:
            target.resume_path = incoming_resume_path
            merged_fields.append("简历")

    if target.supplier_id:
        supplier = db.query(Supplier).filter(Supplier.id == target.supplier_id).first()
        if supplier:
            target.source = supplier.name

    _sync_legacy_fields(target)
    return merged_fields


def _collect_duplicate_matches(data, db: Session) -> List[dict]:
    matches_by_id = {}

    def add_match(candidate: Candidate, reason: str, match_level: str):
        entry = matches_by_id.get(candidate.id)
        if not entry:
            entry = {
                "candidate": candidate,
                "match_reasons": [],
                "match_level": match_level,
            }
            matches_by_id[candidate.id] = entry
        if reason not in entry["match_reasons"]:
            entry["match_reasons"].append(reason)
        if match_level == "exact":
            entry["match_level"] = "exact"

    phone = _normalize_text(getattr(data, "phone", None))
    email = _normalize_text(getattr(data, "email", None))
    name = _normalize_text(getattr(data, "name", None))
    last_company = _normalize_text(getattr(data, "last_company", None))

    if phone:
        candidates = db.query(Candidate).filter(
            Candidate.phone == phone,
            Candidate.deleted_at.is_(None),
        ).all()
        for candidate in candidates:
            add_match(candidate, "手机号相同", "exact")

    if email:
        candidates = db.query(Candidate).filter(
            func.lower(Candidate.email) == email.lower(),
            Candidate.deleted_at.is_(None),
        ).all()
        for candidate in candidates:
            add_match(candidate, "邮箱相同", "exact")

    if name and last_company:
        candidates = db.query(Candidate).filter(
            Candidate.name == name,
            Candidate.last_company == last_company,
            Candidate.deleted_at.is_(None),
        ).all()
        for candidate in candidates:
            add_match(candidate, "姓名和当前公司相同", "fuzzy")

    items = sorted(
        matches_by_id.values(),
        key=lambda item: (
            0 if item["match_level"] == "exact" else 1,
            0 if _active_link_context(item["candidate"].id, db) else 1,
            -item["candidate"].id,
        ),
    )

    results = []
    for item in items:
        candidate = item["candidate"]
        active_link = _active_link_context(candidate.id, db)
        results.append({
            "id": candidate.id,
            "display_id": f"C{candidate.id:03d}",
            "name": candidate.name,
            "phone": candidate.phone,
            "email": candidate.email,
            "last_company": candidate.last_company,
            "last_title": candidate.last_title,
            "match_reasons": item["match_reasons"],
            "match_level": item["match_level"],
            "is_blacklisted": bool(candidate.blacklisted),
            "blacklist_reason": candidate.blacklist_reason,
            "last_application": _last_application_context(candidate.id, db),
            "active_link": active_link,
        })
    return results


@router.post("")
def create_candidate(data: CandidateCreate, db: Session = Depends(get_db)):
    payload = data.model_dump()
    candidate = _create_candidate_record(payload, db)
    db.commit()
    db.refresh(candidate)
    return _candidate_out(candidate)


class DuplicateCheckRequest(BaseModel):
    name: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    last_company: Optional[str] = None


class IntakeResolveRequest(BaseModel):
    decision: Literal["create_new", "merge_existing"]
    incoming: CandidateCreate
    existing_candidate_id: Optional[int] = None
    overwrite_resume: bool = False


@router.post("/check-duplicate")
def check_duplicate(data: DuplicateCheckRequest, db: Session = Depends(get_db)):
    matches = _collect_duplicate_matches(data, db)
    has_blocking = any(item.get("active_link") for item in matches)
    return {
        "matches": matches,
        "requires_decision": bool(matches),
        "has_blocking_in_progress_match": has_blocking,
    }


@router.post("/intake/resolve")
def resolve_intake(data: IntakeResolveRequest, db: Session = Depends(get_db)):
    incoming_payload = data.incoming.model_dump()
    dedup_probe = DuplicateCheckRequest(
        name=incoming_payload.get("name"),
        phone=incoming_payload.get("phone"),
        email=incoming_payload.get("email"),
        last_company=incoming_payload.get("last_company"),
    )
    matches = _collect_duplicate_matches(dedup_probe, db)
    blocking_match = next((item for item in matches if item.get("active_link")), None)

    if data.decision == "create_new":
        if blocking_match:
            active_link = blocking_match.get("active_link") or {}
            target_name = blocking_match.get("name") or blocking_match.get("display_id")
            job_title = active_link.get("job_title") or f"岗位#{active_link.get('job_id')}"
            stage = active_link.get("stage") or "未知阶段"
            raise HTTPException(
                status_code=409,
                detail=f"命中在途候选人「{target_name}」({job_title} · {stage})，请合并后在进行中页处理",
            )
        candidate = _create_candidate_record(incoming_payload, db)
        db.commit()
        db.refresh(candidate)
        return {
            "action": "created",
            "candidate": _candidate_out(candidate),
            "active_link": _active_link_context(candidate.id, db),
        }

    if not data.existing_candidate_id:
        raise HTTPException(status_code=400, detail="合并模式必须选择已有候选人")

    target = db.query(Candidate).filter(
        Candidate.id == data.existing_candidate_id,
        Candidate.deleted_at.is_(None),
    ).first()
    if not target:
        raise HTTPException(status_code=404, detail="目标候选人不存在或已被合并")

    merged_fields = _merge_intake_payload(target, incoming_payload, data.overwrite_resume, db)
    target.updated_at = datetime.utcnow()

    detail_parts = ["导入命中查重：确认同一候选人并合并本次信息"]
    if data.overwrite_resume and _normalize_text(incoming_payload.get("resume_path")):
        detail_parts.append("已覆盖简历")
    if merged_fields:
        detail_parts.append(f"补充字段：{'、'.join(merged_fields)}")

    db.add(HistoryEntry(
        candidate_id=target.id,
        event_type="note",
        detail="；".join(detail_parts),
    ))
    db.commit()
    db.refresh(target)
    return {
        "action": "merged",
        "candidate": _candidate_out(target),
        "active_link": _active_link_context(target.id, db),
        "merged_fields": merged_fields,
    }


@router.get("")
def list_candidates(
    q: Optional[str] = Query(None),
    education: Optional[str] = Query(None),
    tag: Optional[str] = Query(None),
    tags: Optional[str] = Query(None, description="多标签搜索，逗号分隔"),
    tag_mode: Optional[str] = Query("all", description="标签匹配模式: all/any"),
    followup_status: Optional[str] = Query(None),
    source: Optional[str] = Query(None),
    starred: Optional[bool] = Query(None),
    show_blacklisted: Optional[bool] = Query(False),
    db: Session = Depends(get_db),
):
    query = db.query(Candidate).filter(Candidate.deleted_at.is_(None))
    if not show_blacklisted:
        query = query.filter(Candidate.blacklisted == False)
    if q:
        query = query.filter(
            or_(
                Candidate.name.ilike(f"%{q}%"),
                Candidate.name_en.ilike(f"%{q}%"),
                Candidate.phone.ilike(f"%{q}%"),
                Candidate.email.ilike(f"%{q}%"),
            )
        )
    if education:
        query = query.filter(Candidate.education.ilike(f"%{education}%"))
    if followup_status:
        query = query.filter(Candidate.followup_status == followup_status)
    if source:
        query = query.filter(Candidate.source == source)
    if starred is True:
        query = query.filter(Candidate.starred == 1)
    candidates = query.order_by(Candidate.created_at.desc()).all()
    required_tags = _parse_tag_query(tags)
    if tag:
        required_tags = _normalize_tags(required_tags + [tag])
    if required_tags:
        mode = (tag_mode or "all").lower()
        if mode not in {"all", "any"}:
            mode = "all"
        required_set = set(t.lower() for t in required_tags)
        if mode == "any":
            candidates = [
                c for c in candidates
                if required_set.intersection(set(x.lower() for x in _normalize_tags(c.skill_tags or [])))
            ]
        else:
            candidates = [
                c for c in candidates
                if required_set.issubset(set(x.lower() for x in _normalize_tags(c.skill_tags or [])))
            ]
    result = []
    for c in candidates:
        d = _candidate_out(c)
        d["active_links"] = [
            {"job_id": lnk.job_id, "job_title": lnk.job.title if lnk.job else None, "stage": lnk.stage}
            for lnk in c.job_links if not lnk.outcome
        ]
        d["job_links"] = [
            {"id": lnk.id, "job_id": lnk.job_id, "job_title": lnk.job.title if lnk.job else None, "outcome": lnk.outcome}
            for lnk in c.job_links
        ]
        result.append(d)
    return result


@router.get("/{candidate_id}")
def get_candidate(candidate_id: int, db: Session = Depends(get_db)):
    c = db.query(Candidate).filter(Candidate.id == candidate_id, Candidate.deleted_at.is_(None)).first()
    if not c:
        raise HTTPException(status_code=404, detail="候选人不存在或已被合并")
    result = _candidate_out(c)
    result["history"] = [
        {
            "id": h.id,
            "event_type": h.event_type,
            "detail": h.detail,
            "job_id": h.job_id,
            "timestamp": h.timestamp.isoformat() if h.timestamp else None,
        }
        for h in sorted(c.history, key=lambda x: x.timestamp or datetime.min, reverse=True)
    ]
    result["job_links"] = [
        {
            "id": lnk.id,
            "job_id": lnk.job_id,
            "job_title": lnk.job.title if lnk.job else None,
            "stage": lnk.stage,
            "notes": lnk.notes,
            "outcome": lnk.outcome,
            "rejection_reason": lnk.rejection_reason,
            "created_at": lnk.created_at.isoformat() if lnk.created_at else None,
        }
        for lnk in c.job_links
    ]
    return result


@router.patch("/{candidate_id}")
def update_candidate(candidate_id: int, data: CandidateUpdate, db: Session = Depends(get_db)):
    c = db.query(Candidate).filter(Candidate.id == candidate_id, Candidate.deleted_at.is_(None)).first()
    if not c:
        raise HTTPException(status_code=404, detail="候选人不存在")
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(c, field, value)
    # Auto-fill source from supplier name
    if c.supplier_id:
        supplier = db.query(Supplier).filter(Supplier.id == c.supplier_id).first()
        if supplier:
            c.source = supplier.name
    _sync_legacy_fields(c)
    c.updated_at = datetime.utcnow()
    db.add(HistoryEntry(candidate_id=c.id, event_type="updated", detail="信息已更新"))
    db.commit()
    db.refresh(c)
    return _candidate_out(c)


class BlacklistRequest(BaseModel):
    reason: str
    note: Optional[str] = None


@router.post("/{candidate_id}/blacklist")
def blacklist_candidate(candidate_id: int, data: BlacklistRequest, db: Session = Depends(get_db)):
    c = db.query(Candidate).filter(Candidate.id == candidate_id, Candidate.deleted_at.is_(None)).first()
    if not c:
        raise HTTPException(status_code=404, detail="候选人不存在")
    c = candidates_svc.blacklist_candidate(db, c, data.reason, data.note)
    return _candidate_out(c)


class UnblacklistRequest(BaseModel):
    reason: str


@router.delete("/{candidate_id}/blacklist")
def unblacklist_candidate(candidate_id: int, data: UnblacklistRequest, db: Session = Depends(get_db)):
    c = db.query(Candidate).filter(Candidate.id == candidate_id, Candidate.deleted_at.is_(None)).first()
    if not c:
        raise HTTPException(status_code=404, detail="候选人不存在")
    c = candidates_svc.unblacklist_candidate(db, c, data.reason)
    return _candidate_out(c)


@router.get("/{candidate_id}/last-application")
def get_last_application(candidate_id: int, db: Session = Depends(get_db)):
    c = db.query(Candidate).filter(Candidate.id == candidate_id, Candidate.deleted_at.is_(None)).first()
    if not c:
        raise HTTPException(status_code=404, detail="候选人不存在")
    return {"last_application": _last_application_context(candidate_id, db)}


@router.get("/{candidate_id}/resume-preview")
def resume_preview(candidate_id: int, db: Session = Depends(get_db)):
    c = db.query(Candidate).filter(Candidate.id == candidate_id, Candidate.deleted_at.is_(None)).first()
    if not c or not c.resume_path:
        raise HTTPException(status_code=404, detail="简历不存在")
    if not os.path.exists(c.resume_path):
        raise HTTPException(status_code=404, detail="简历文件不存在")
    ext = os.path.splitext(c.resume_path)[1].lower()
    # 构建静态访问路径
    parts = c.resume_path.replace("\\", "/").split("/")
    static_path = "/resumes/" + "/".join(parts[-2:])
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(c.resume_path)
            html_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    html_parts.append(f"<p>{escape(text)}</p>")
            for table in doc.tables:
                html_parts.append("<table style='border-collapse:collapse;width:100%'>")
                for row in table.rows:
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        html_parts.append(
                            "<td style='border:1px solid #ddd;padding:6px 8px;font-size:13px'>"
                            f"{escape(cell.text)}</td>"
                        )
                    html_parts.append("</tr>")
                html_parts.append("</table>")
            return {"html": "".join(html_parts)}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX 解析失败: {str(e)}")
    return {"redirect": static_path}
